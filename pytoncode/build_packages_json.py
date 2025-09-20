# -*- coding: utf-8 -*-
"""
يبني ملف JSON بالهيكل المطلوب اعتماداً على:
- حدود.docx  → ترتيب الحزم/التصنيفات/البوتات + نص الحدود
- نبذة.docx  → نص النبذة لكل بوت (يدعم @@@عنوان ثم النص، أو "عنوان":"النص")
- مثال.docx  → نص المثال لكل بوت (يدعم "العنوان":"النص" في نفس السطر،
               أو كسر سطر مع 'الوصف (مثال): "النص"')
- روابط النسخة الكاملة.docx → روابط النماذج 4O/5 (يدعم الجداول والروابط داخل الخلايا)

احفظه باسم build_packages_json.py ثم شغّله.
"""

import json
import os
import re
import difflib
from collections import OrderedDict, defaultdict
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

# ======== إعدادات المسارات ========
# إن كانت ملفاتك في /mnt/data كما في جلسة العمل الحالية، اترك BASE كما هو.
BASE = Path(__file__).resolve().parent
HUDUD_PATH = os.path.join(BASE, "حدود.docx")
NOBTHA_PATH = os.path.join(BASE, "نبذة.docx")
MITHAL_PATH = os.path.join(BASE, "مثال.docx")
# LINKS_PATH  = os.path.join(BASE, "روابط النسخة الكاملة.docx")
LINKS_PATH  = os.path.join(BASE, "old.docx")
OUTPUT_JSON = os.path.join(BASE, "output.json")


# ======== أدوات مساعدة ========
AR_QUOTE_CHARS = '«»“”„‟‚‛'
EN_QUOTE_CHARS = '"\''
PUNCT_TO_STRIP = '：:؛،,'

def normalize_title(s: str) -> str:
    """تطبيع اسم البوت/التصنيف بإزالة محارف البداية والنهاية والاقتباسات."""
    if not s:
        return ""
    s = s.strip()
    s = re.sub(r'^\s*#\s*', '', s)  # أزل بادئة #
    s = s.strip(AR_QUOTE_CHARS + EN_QUOTE_CHARS + PUNCT_TO_STRIP)
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def norm_for_match(s: str) -> str:
    """تطبيع أخف لاستخدامه في المطابقة التقريبية."""
    if not s: return ""
    s = re.sub(r'[\u200f\u200e]', '', s)  # محارف اتجاهية
    s = s.replace("ـ", "")                # كَشْدَة ممتدة
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def best_match_title(text, known_titles, cutoff=0.88):
    """يعيد أفضل عنوان معروف يظهر داخل النص أو أقربه تقريبياً."""
    text_n = norm_for_match(text)
    # 1) احتواء مباشر (نختار الأطول)
    cands = [t for t in known_titles if t and norm_for_match(t) in text_n]
    if cands:
        return max(cands, key=len)
    # 2) مطابقة تقريبية
    matches = difflib.get_close_matches(text_n, [norm_for_match(t) for t in known_titles], n=1, cutoff=cutoff)
    if matches:
        # أعد العنوان الأصلي المطابق للمطابقة المُطبّعة
        match_n = matches[0]
        for t in known_titles:
            if norm_for_match(t) == match_n:
                return t
    return None

def read_docx_lines(path: str):
    """قراءة جميع الفقرات غير الفارغة كسطور نصية."""
    if not os.path.exists(path):
        return []
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    return lines

def is_package_line(line: str):
    # مثال: "باقة ...."
    return bool(re.match(r'^\s*باقة\s+', line))

def is_category_line(line: str):
    # يدعم: "تصنيف ..." أو "نماذج ..."
    return bool(re.match(r'^\s*(تصنيف|نماذج)\s+', line))

def is_bot_header_line(line: str):
    # سطر يبدأ بـ #عنوان_البوت
    return bool(re.match(r'^\s*#', line))


# ======== تحليل حدود.docx → (حِزم ← تصنيفات ← بوتات + نص الحدود) ========
def parse_hudud(lines):
    """
    يُعيد:
    OrderedDict {
        package_name: {'categories': OrderedDict {
            category_name: {'bots': OrderedDict {
                bot_title: {'hudud': '...'}
            }}
        }}
    }
    """
    packages = OrderedDict()
    current_package = None
    current_category = None
    current_bot = None
    buffer = []

    def flush_bot():
        nonlocal buffer, current_package, current_category, current_bot
        if current_package and current_category and current_bot:
            text = "\n".join(buffer).strip()
            packages.setdefault(current_package, {'categories': OrderedDict()})
            cats = packages[current_package]['categories']
            cats.setdefault(current_category, {'bots': OrderedDict()})
            bots = cats[current_category]['bots']
            bots.setdefault(current_bot, {})
            bots[current_bot]['hudud'] = text
        buffer = []

    for line in lines:
        if is_package_line(line):
            flush_bot()
            current_package = normalize_title(line)
            current_category = None
            current_bot = None
        elif is_category_line(line):
            flush_bot()
            current_category = normalize_title(line)
            current_bot = None
        elif is_bot_header_line(line):
            flush_bot()
            current_bot = normalize_title(line)
        else:
            if current_bot:
                buffer.append(line)
    flush_bot()
    return packages


# ======== تحليل نبذة.docx → {bot_title: 'نبذة...'} ========
def parse_nobtha(lines):
    """
    يدعم شكلين:
      1) @@@عنوان_البوت
         ...النص حتى @@@ التالي
      2) "العنوان":"النص"  (بنفس السطر أو بصيغة كسطرين مع "الوصف (نبذة):")
    """
    # أولاً: محاولة نمط الأزواج (مفيد لو الملف منظّم بهذه الطريقة)
    text = "\n".join(lines)
    pair_pattern = re.compile(
        r'[\"“](.+?)[\"”]\s*[:：]\s*'                 # "العنوان":
        r'(?:\r?\n\s*)?'                              # كسر سطر اختياري
        r'(?:الوصف\s*\(?\s*نبذة\s*\)?\s*[:：]\s*)?'   # "الوصف (نبذة):" اختياري
        r'[\"“](.+?)[\"”]',                           # "النص"
        flags=re.MULTILINE
    )
    result = {}
    for m in pair_pattern.finditer(text):
        title = normalize_title(m.group(1))
        desc  = m.group(2).strip()
        result[title] = desc

    # ثانياً: دعم أسلوب @@@
    current_title = None
    buffer = []
    def flush():
        nonlocal current_title, buffer
        if current_title and buffer and current_title not in result:
            result[current_title] = "\n".join(buffer).strip()
        buffer = []

    for line in lines:
        if line.strip().startswith('@@@'):
            flush()
            current_title = normalize_title(line.replace('@@@', '', 1))
        else:
            if current_title:
                buffer.append(line)
    flush()
    return result


# ======== تحليل مثال.docx → {bot_title: 'مثال...'} ========
def parse_mithal(lines):
    """
    يلتقط الصيغ:
    1) "العنوان": "النص"
    2) "العنوان":
       الوصف (مثال): "النص"
    """
    text = "\n".join(lines)
    pair_pattern = re.compile(
        r'[\"“](.+?)[\"”]\s*[:：]\s*'               # "العنوان":
        r'(?:\r?\n\s*)?'                            # كسر سطر اختياري
        r'(?:الوصف\s*\(?\s*مثال\s*\)?\s*[:：]\s*)?' # "الوصف (مثال):" اختياري
        r'[\"“](.+?)[\"”]',                         # "النص"
        flags=re.MULTILINE
    )

    result = {}
    for m in pair_pattern.finditer(text):
        title = normalize_title(m.group(1))
        desc  = m.group(2).strip()
        result[title] = desc
    return result


# ======== تحليل روابط النسخة الكاملة.docx (يدعم الجداول والروابط) ========
def parse_links(path, known_titles):
    """
    يقرأ الروابط من الجداول والفقرات، ويُرجع:
      { 'عنوان البوت': {'4O': url_or_empty, '5': url_or_empty}, ... }
    """
    if not os.path.exists(path):
        return {}

    doc = Document(path)
    known_titles = list(known_titles)  # لضمان قابلية الفهرسة

    def clean_title_in_cell(s: str) -> str:
        if not s: return ""
        s = s.replace("🔗", "").strip()
        s = re.sub(r'\s*[–\-—]\s*نموذج\s*[45](?:o|O)?\s*$', '', s).strip()
        s = re.sub(r'\s+', ' ', s).strip()
        return s

    def detect_model_from_text(txt: str):
        m = re.search(r'نموذج\s*([45](?:o|O)?)', txt, re.IGNORECASE)
        if not m: return None
        g = m.group(1)
        if g in ('5', '٥'): return '5'
        return '4O'

    def detect_model_from_url(url: str):
        u = url.lower()
        if any(k in u for k in ['mod-5', '/5', 'gpt-5', 'model-5']):
            return '5'
        if any(k in u for k in ['mod-4o', 'gpt-4o', '4o', 'model-4o', 'mod-4']):
            return '4O'
        return None

    def hyperlinks_in_paragraph(p):
        # يعيد قائمة r:id الموجودة داخل الفقرة
        return [hl.get(qn('r:id'))
                for hl in p._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
                if hl.get(qn('r:id'))]

    result = defaultdict(lambda: {'4O': '', '5': ''})

    # --- 1) الجداول: نختار عنوان الصف ثم نربط كل روابط الصف به ---
    for table in doc.tables:
        for row in table.rows:
            # اجمع نصوص الصف بالكامل لتحديد العنوان الأفضل
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            row_title = best_match_title(clean_title_in_cell(row_text), known_titles)
            # اجمع كل الروابط في الصف
            for cell in row.cells:
                cell_text = cell.text.strip()
                # إن لم يوجد عنوان على مستوى الصف، جرّب على مستوى الخلية
                cell_title = row_title or best_match_title(clean_title_in_cell(cell_text), known_titles)

                rids = []
                for p in cell.paragraphs:
                    rids.extend(hyperlinks_in_paragraph(p))

                for rId in rids:
                    rel = doc.part.rels.get(rId)
                    if not rel:
                        continue
                    url = rel.target_ref
                    model = detect_model_from_text(cell_text) or detect_model_from_url(url)
                    # لو لم نعرف الموديل، عيّنه مؤقتًا 4O لتعبئة خانة واحدة على الأقل
                    if not model:
                        model = '4O'
                    title_key = cell_title
                    if title_key:
                        result[title_key][model] = url

    # --- 2) الفقرات الحرة خارج الجداول (شبكة أمان) ---
    for p in doc.paragraphs:
        rids = hyperlinks_in_paragraph(p)
        if not rids:
            continue
        t = p.text.strip()
        title_in_p = best_match_title(clean_title_in_cell(t), known_titles)
        for rId in rids:
            rel = doc.part.rels.get(rId)
            if not rel:
                continue
            url = rel.target_ref
            model = detect_model_from_text(t) or detect_model_from_url(url) or '4O'
            if title_in_p:
                result[title_in_p][model] = url

    return dict(result)


# ======== بناء JSON النهائي ========
def build_json_from_docs():
    hudud_lines = read_docx_lines(HUDUD_PATH)
    nobtha_lines = read_docx_lines(NOBTHA_PATH)
    mithal_lines = read_docx_lines(MITHAL_PATH)

    packages = parse_hudud(hudud_lines)
    nobtha_map = parse_nobtha(nobtha_lines)
    mithal_map = parse_mithal(mithal_lines)

    # قائمة بكل عناوين البوتات المعروفة من حدود.docx
    known_bot_titles = []
    for pkg in packages.values():
        for cat in pkg['categories'].values():
            known_bot_titles.extend(list(cat['bots'].keys()))

    links_map = parse_links(LINKS_PATH, known_bot_titles)

    # تحويل التركيب إلى الشكل النهائي
    out = {"packages": []}
    package_id_counter = 1

    for package_name, pkg_obj in packages.items():
        package_entry = {
            "package": package_name,
            "packageId": package_id_counter,
            "categories": []
        }
        package_id_counter += 1

        for category_name, cat_obj in pkg_obj['categories'].items():
            category_entry = {
                "category": category_name,
                "bots": []
            }
            for bot_title, bot_obj in cat_obj['bots'].items():
                bot_entry = {
                    "botTitle": bot_title,
                    "النموذج": {
                        "4O": links_map.get(bot_title, {}).get("4O", ""),
                        "5":  links_map.get(bot_title, {}).get("5", "")
                    },
                    "نبذة": nobtha_map.get(bot_title, ""),
                    "حدود": bot_obj.get("hudud", ""),
                    "مثال": mithal_map.get(bot_title, "")
                }
                category_entry["bots"].append(bot_entry)

            package_entry["categories"].append(category_entry)
        out["packages"].append(package_entry)

    # حفظ JSON
    with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)

    return out


if __name__ == "__main__":
    data = build_json_from_docs()
    print(f"✅ تم إنشاء الملف: {OUTPUT_JSON}")
    # ملخص سريع
    for p in data["packages"]:
        print(f"- {p['package']} (id={p['packageId']}): {len(p['categories'])} تصنيف")
        for c in p["categories"]:
            print(f"  • {c['category']}: {len(c['bots'])} بوت")
